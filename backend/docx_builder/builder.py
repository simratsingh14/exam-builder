"""DOCX generation engine for exam papers and answer keys.

Converts a ``Paper`` model into a formatted ``.docx`` document using
python-docx. TipTap JSON content is traversed recursively and mapped to
the appropriate python-docx primitives.
"""

import os
from io import BytesIO
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from ..models import (
    ImageQuestion,
    MCQQuestion,
    Paper,
    PaperHeader,
    PaperStyle,
    Question,
    TableQuestion,
    TextQuestion,
)


# ── Data-dir helper (read at call-time so tests can monkeypatch) ──────────────


def _data_dir() -> Path:
    return Path(os.getenv("DATA_DIR", "/data"))


# ── Public API ────────────────────────────────────────────────────────────────


def build_docx(paper: Paper) -> bytes:
    """Build a complete exam paper ``.docx`` from a ``Paper`` model.

    Args:
        paper: Fully populated exam paper domain object.

    Returns:
        Raw bytes of a valid ``.docx`` file.
    """
    doc = Document()
    _apply_margins(doc, paper.style)
    _apply_default_font(doc, paper.style)
    _add_header_footer(doc, paper.style)
    _add_paper_header(doc, paper.header, paper.style)
    _add_questions(doc, paper.questions, paper.style)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_answer_key(paper: Paper) -> bytes:
    """Build an answer-key ``.docx`` listing correct MCQ answers only.

    Args:
        paper: Paper whose MCQ questions have ``is_correct`` flags set.

    Returns:
        Raw bytes of a valid ``.docx`` file.
    """
    doc = Document()
    _apply_margins(doc, paper.style)
    _apply_default_font(doc, paper.style)

    title = paper.header.title or "Exam"
    doc.add_heading(f"Answer Key: {title}", 0)

    mcq_num = 1
    for q in paper.questions:
        if q.type == "mcq":
            assert isinstance(q, MCQQuestion)
            correct = next((opt.label for opt in q.options if opt.is_correct), "N/A")
            doc.add_paragraph(f"Q{mcq_num}: {correct}")
            mcq_num += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Document-level helpers ────────────────────────────────────────────────────


def _apply_margins(doc: Document, style: PaperStyle) -> None:
    """Set page margins on every section of the document."""
    for section in doc.sections:
        section.top_margin = Inches(style.margin_top)
        section.bottom_margin = Inches(style.margin_bottom)
        section.left_margin = Inches(style.margin_left)
        section.right_margin = Inches(style.margin_right)


def _apply_default_font(doc: Document, style: PaperStyle) -> None:
    """Apply font family and size to the Normal paragraph style."""
    normal = doc.styles["Normal"]
    normal.font.name = style.font_family
    normal.font.size = Pt(style.font_size)


def _add_header_footer(doc: Document, style: PaperStyle) -> None:
    """Write header and footer text into the first section."""
    section = doc.sections[0]
    if style.header_text:
        para = section.header.paragraphs[0]
        para.text = style.header_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if style.footer_text:
        para = section.footer.paragraphs[0]
        para.text = style.footer_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_paper_header(doc: Document, header: PaperHeader, style: PaperStyle) -> None:
    """Render the institutional header block (logo, title, details row, rule)."""
    # Logo
    if style.logo_filename:
        logo_path = _data_dir() / "uploads" / style.logo_filename
        if logo_path.exists():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(str(logo_path), width=Inches(1.5))

    # Institution name
    if header.institution:
        para = doc.add_paragraph(header.institution)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(style.font_size + 4)

    # Exam title
    if header.title:
        para = doc.add_paragraph(header.title)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(style.font_size + 2)

    # Details row: subject | date | duration | total marks
    details: list[str] = []
    if header.subject:
        details.append(f"Subject: {header.subject}")
    if header.date:
        details.append(f"Date: {header.date}")
    if header.duration:
        details.append(f"Duration: {header.duration}")
    if header.total_marks:
        # Format as integer when value is whole number (e.g. 100.0 → "100")
        marks_str = f"{int(header.total_marks)}" if header.total_marks == int(header.total_marks) else f"{header.total_marks}"
        details.append(f"Total Marks: {marks_str}")
    if details:
        para = doc.add_paragraph("   |   ".join(details))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Horizontal rule beneath the header block
    _add_horizontal_rule(doc)
    doc.add_paragraph()  # visual spacer


def _add_horizontal_rule(doc: Document) -> None:
    """Insert a paragraph with a bottom border (acts as a horizontal rule)."""
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ── Question rendering ────────────────────────────────────────────────────────


def _add_questions(doc: Document, questions: Sequence[Question], style: PaperStyle) -> None:
    """Render all questions, inserting section headings when the section changes."""
    current_section: str = ""
    for num, q in enumerate(questions, start=1):
        # Section heading when section label changes
        if q.section and q.section != current_section:
            current_section = q.section
            doc.add_heading(current_section, level=2)

        marks_str = _marks_label(q.marks)

        if isinstance(q, TextQuestion):
            _write_question_prefix(doc, num, marks_str)
            _tiptap_to_doc(doc, q.content, style)

        elif isinstance(q, MCQQuestion):
            _write_question_prefix(doc, num, marks_str)
            _tiptap_to_doc(doc, q.stem, style)
            for opt in q.options:
                doc.add_paragraph(f"    ({opt.label}) {opt.text}")

        elif isinstance(q, TableQuestion):
            _write_question_prefix(doc, num, marks_str)
            _tiptap_to_doc(doc, q.content, style)

        elif isinstance(q, ImageQuestion):
            _write_question_prefix(doc, num, marks_str)
            img_path = _data_dir() / "uploads" / q.filename
            if img_path.exists():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(str(img_path), width=Inches(4))
            if q.caption:
                cap = doc.add_paragraph(q.caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cap.runs:
                    cap.runs[0].italic = True

        doc.add_paragraph()  # spacer between questions


def _marks_label(marks: float) -> str:
    """Return a formatted marks label, or empty string if marks is zero."""
    if not marks:
        return ""
    count = int(marks) if marks == int(marks) else marks
    return f"[{count} mark{'s' if marks != 1 else ''}]"


def _write_question_prefix(doc: Document, num: int, marks_str: str) -> None:
    """Add a bold question number paragraph with optional italic marks label."""
    para = doc.add_paragraph()
    run = para.add_run(f"Q{num}.")
    run.bold = True
    if marks_str:
        run2 = para.add_run(f"  {marks_str}")
        run2.italic = True


# ── TipTap JSON → python-docx ─────────────────────────────────────────────────


def _tiptap_text(node: dict) -> str:
    """Recursively extract plain text from a TipTap JSON node."""
    if node.get("type") == "text":
        return node.get("text", "")
    return "".join(_tiptap_text(child) for child in node.get("content", []))


def _tiptap_to_doc(doc: Document, node: dict, style: PaperStyle) -> None:
    """Recursively map TipTap JSON nodes to python-docx document elements.

    Handles: doc, paragraph (with text marks), heading, bulletList,
    orderedList, table/tableRow/tableCell, and inline images.
    """
    node_type = node.get("type", "")

    if node_type == "doc":
        for child in node.get("content", []):
            _tiptap_to_doc(doc, child, style)

    elif node_type == "paragraph":
        para = doc.add_paragraph()
        for child in node.get("content", []):
            child_type = child.get("type", "")
            if child_type == "text":
                run = para.add_run(child.get("text", ""))
                for mark in child.get("marks", []):
                    mark_type = mark.get("type", "")
                    if mark_type == "bold":
                        run.bold = True
                    elif mark_type == "italic":
                        run.italic = True
                    elif mark_type == "underline":
                        run.underline = True
            elif child_type == "image":
                src: str = child.get("attrs", {}).get("src", "")
                if "/api/uploads/" in src:
                    fname = src.rsplit("/", 1)[-1]
                    img_path = _data_dir() / "uploads" / fname
                    if img_path.exists():
                        para.add_run().add_picture(str(img_path), width=Inches(3))

    elif node_type == "heading":
        level: int = node.get("attrs", {}).get("level", 1)
        doc.add_heading(_tiptap_text(node), level=level)

    elif node_type == "bulletList":
        for item in node.get("content", []):
            doc.add_paragraph(_tiptap_text(item), style="List Bullet")

    elif node_type == "orderedList":
        for item in node.get("content", []):
            doc.add_paragraph(_tiptap_text(item), style="List Number")

    elif node_type == "table":
        rows = node.get("content", [])  # tableRow nodes
        if not rows:
            return
        n_rows = len(rows)
        n_cols = max(len(row.get("content", [])) for row in rows)
        if n_cols == 0:
            return
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell_node in enumerate(row.get("content", [])):
                table.cell(i, j).text = _tiptap_text(cell_node)
