"""Tests for the DOCX generation engine.

Written before implementation (TDD red phase).
Tests drive the builder to handle all question types,
styling, and edge cases.
"""

import struct
import zlib
from io import BytesIO

import pytest
from docx import Document as DocxDocument

from backend.models import (
    ImageQuestion,
    MCQOption,
    MCQQuestion,
    Paper,
    PaperHeader,
    PaperStyle,
    TableQuestion,
    TextQuestion,
)
from backend.docx_builder.builder import build_docx, build_answer_key


# ── Helpers ───────────────────────────────────────────────────────────────────


def _open(docx_bytes: bytes) -> DocxDocument:
    return DocxDocument(BytesIO(docx_bytes))


def _full_text(doc: DocxDocument) -> str:
    return " ".join(p.text for p in doc.paragraphs)


def _para_text(doc: DocxDocument) -> list[str]:
    return [p.text for p in doc.paragraphs if p.text.strip()]


def _tiptap_para(text: str) -> dict:
    """Build a minimal TipTap document with a single paragraph."""
    return {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": text}],
            }
        ],
    }


def _tiptap_empty() -> dict:
    return {"type": "doc", "content": []}


# ── Core ──────────────────────────────────────────────────────────────────────


def test_build_docx_returns_bytes() -> None:
    paper = Paper()
    result = build_docx(paper)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_empty_paper_produces_valid_docx() -> None:
    paper = Paper()
    doc = _open(build_docx(paper))
    assert doc is not None


def test_docx_contains_exam_title() -> None:
    paper = Paper(header=PaperHeader(title="Science Midterm"))
    doc = _open(build_docx(paper))
    assert "Science Midterm" in _full_text(doc)


def test_docx_contains_institution() -> None:
    paper = Paper(header=PaperHeader(institution="Springfield High"))
    doc = _open(build_docx(paper))
    assert "Springfield High" in _full_text(doc)


def test_docx_contains_subject_and_date() -> None:
    paper = Paper(header=PaperHeader(subject="Biology", date="2026-03-01"))
    doc = _open(build_docx(paper))
    text = _full_text(doc)
    assert "Biology" in text
    assert "2026-03-01" in text


def test_docx_contains_duration_and_total_marks() -> None:
    paper = Paper(header=PaperHeader(duration="2 hours", total_marks=100))
    doc = _open(build_docx(paper))
    text = _full_text(doc)
    assert "2 hours" in text
    assert "100" in text


# ── Text question ─────────────────────────────────────────────────────────────


def test_text_question_appears_in_docx() -> None:
    paper = Paper(
        questions=[TextQuestion(content=_tiptap_para("What is photosynthesis?"))]
    )
    assert "What is photosynthesis?" in _full_text(_open(build_docx(paper)))


def test_text_question_shows_marks() -> None:
    paper = Paper(
        questions=[TextQuestion(content=_tiptap_para("Explain gravity."), marks=5)]
    )
    assert "5" in _full_text(_open(build_docx(paper)))


def test_text_question_shows_section_heading() -> None:
    paper = Paper(
        questions=[TextQuestion(content=_tiptap_para("Q text."), section="Section A")]
    )
    assert "Section A" in _full_text(_open(build_docx(paper)))


def test_multiple_text_questions_numbered() -> None:
    paper = Paper(
        questions=[
            TextQuestion(content=_tiptap_para("First question.")),
            TextQuestion(content=_tiptap_para("Second question.")),
        ]
    )
    text = _full_text(_open(build_docx(paper)))
    assert "Q1" in text
    assert "Q2" in text


# ── MCQ question ──────────────────────────────────────────────────────────────


def test_mcq_stem_and_options_appear() -> None:
    paper = Paper(
        questions=[
            MCQQuestion(
                stem=_tiptap_para("What is 2 + 2?"),
                options=[
                    MCQOption(label="A", text="3"),
                    MCQOption(label="B", text="4", is_correct=True),
                    MCQOption(label="C", text="5"),
                    MCQOption(label="D", text="6"),
                ],
            )
        ]
    )
    text = _full_text(_open(build_docx(paper)))
    assert "What is 2 + 2?" in text
    assert "(A)" in text or "A)" in text
    assert "(B)" in text or "B)" in text
    assert "4" in text


def test_mcq_correct_answer_not_revealed_in_paper() -> None:
    """The paper docx must NOT expose which option is correct."""
    paper = Paper(
        questions=[
            MCQQuestion(
                stem=_tiptap_para("Capital of France?"),
                options=[
                    MCQOption(label="A", text="London"),
                    MCQOption(label="B", text="Paris", is_correct=True),
                    MCQOption(label="C", text="Berlin"),
                    MCQOption(label="D", text="Rome"),
                ],
            )
        ]
    )
    text = _full_text(_open(build_docx(paper)))
    # The word "correct" or "is_correct" must not appear
    assert "is_correct" not in text.lower()
    assert "correct answer" not in text.lower()


# ── Table question ────────────────────────────────────────────────────────────


def test_table_question_produces_table_in_docx() -> None:
    table_content = {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            {"type": "tableCell", "content": [{"type": "paragraph",
                                "content": [{"type": "text", "text": "Header 1"}]}]},
                            {"type": "tableCell", "content": [{"type": "paragraph",
                                "content": [{"type": "text", "text": "Header 2"}]}]},
                        ],
                    },
                    {
                        "type": "tableRow",
                        "content": [
                            {"type": "tableCell", "content": [{"type": "paragraph",
                                "content": [{"type": "text", "text": "Cell A"}]}]},
                            {"type": "tableCell", "content": [{"type": "paragraph",
                                "content": [{"type": "text", "text": "Cell B"}]}]},
                        ],
                    },
                ],
            }
        ],
    }
    paper = Paper(questions=[TableQuestion(content=table_content)])
    doc = _open(build_docx(paper))
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Header 1"
    assert doc.tables[0].cell(1, 1).text == "Cell B"


# ── Image question ────────────────────────────────────────────────────────────


def test_image_question_with_missing_file_does_not_crash() -> None:
    """A missing upload file should be silently skipped, not crash."""
    paper = Paper(
        questions=[ImageQuestion(filename="missing.png", caption="A diagram")]
    )
    result = build_docx(paper)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_image_question_with_real_file(tmp_path, monkeypatch) -> None:
    """When the image exists on disk it should appear in the document."""
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    uploads = tmp_path / "uploads"
    uploads.mkdir(parents=True, exist_ok=True)

    # Write a minimal 1×1 PNG
    def _png() -> bytes:
        def chunk(name: bytes, data: bytes) -> bytes:
            return (struct.pack(">I", len(data)) + name + data
                    + struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF))
        return (b"\x89PNG\r\n\x1a\n"
                + chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
                + chunk(b"IDAT", zlib.compress(b"\x00\xff\xff\xff"))
                + chunk(b"IEND", b""))

    img_path = uploads / "test.png"
    img_path.write_bytes(_png())

    import importlib, backend.docx_builder.builder as b
    importlib.reload(b)

    paper = Paper(questions=[ImageQuestion(filename="test.png", caption="Fig 1")])
    doc = _open(b.build_docx(paper))
    # Caption should appear
    assert "Fig 1" in _full_text(doc)


# ── Answer key ────────────────────────────────────────────────────────────────


def test_answer_key_contains_correct_answers() -> None:
    paper = Paper(
        header=PaperHeader(title="Quiz"),
        questions=[
            MCQQuestion(
                stem=_tiptap_para("Q1"),
                options=[
                    MCQOption(label="A", text="Wrong"),
                    MCQOption(label="B", text="Right", is_correct=True),
                    MCQOption(label="C", text="Wrong"),
                    MCQOption(label="D", text="Wrong"),
                ],
            ),
            MCQQuestion(
                stem=_tiptap_para("Q2"),
                options=[
                    MCQOption(label="A", text="Correct", is_correct=True),
                    MCQOption(label="B", text="Wrong"),
                    MCQOption(label="C", text="Wrong"),
                    MCQOption(label="D", text="Wrong"),
                ],
            ),
        ],
    )
    doc = _open(build_answer_key(paper))
    text = _full_text(doc)
    assert "Q1: B" in text
    assert "Q2: A" in text


def test_answer_key_skips_non_mcq_questions() -> None:
    """Answer key only lists MCQ questions."""
    paper = Paper(
        questions=[
            TextQuestion(content=_tiptap_para("Describe evolution.")),
            MCQQuestion(
                stem=_tiptap_para("MCQ"),
                options=[MCQOption(label="A", text="Yes", is_correct=True),
                         MCQOption(label="B", text="No"),
                         MCQOption(label="C", text="Maybe"),
                         MCQOption(label="D", text="Never")],
            ),
        ]
    )
    doc = _open(build_answer_key(paper))
    text = _full_text(doc)
    assert "Q1: A" in text
    assert "Describe evolution" not in text


def test_answer_key_returns_bytes() -> None:
    paper = Paper(questions=[
        MCQQuestion(stem=_tiptap_para("x"),
                    options=[MCQOption(label="A", text="y", is_correct=True),
                             MCQOption(label="B", text="z"),
                             MCQOption(label="C", text="w"),
                             MCQOption(label="D", text="v")])
    ])
    result = build_answer_key(paper)
    assert isinstance(result, bytes)
    assert len(result) > 0


# ── Styling ───────────────────────────────────────────────────────────────────


def test_page_margins_applied() -> None:
    from docx.shared import Inches
    style = PaperStyle(margin_top=1.5, margin_bottom=1.5, margin_left=2.0, margin_right=2.0)
    paper = Paper(style=style)
    doc = _open(build_docx(paper))
    section = doc.sections[0]
    assert abs(section.top_margin - Inches(1.5)) < 100  # EMU tolerance
    assert abs(section.left_margin - Inches(2.0)) < 100


def test_header_text_applied() -> None:
    style = PaperStyle(header_text="CONFIDENTIAL")
    paper = Paper(style=style)
    doc = _open(build_docx(paper))
    header_text = " ".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "CONFIDENTIAL" in header_text


def test_footer_text_applied() -> None:
    style = PaperStyle(footer_text="Page 1")
    paper = Paper(style=style)
    doc = _open(build_docx(paper))
    footer_text = " ".join(p.text for p in doc.sections[0].footer.paragraphs)
    assert "Page 1" in footer_text
